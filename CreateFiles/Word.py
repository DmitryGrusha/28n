from docx import Document
from docx.shared import Inches, Pt
from docx.enum.section import WD_ORIENTATION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

from Logic.GlobalHandlers import *


def createWordA5(fullName: str, sex: str, birthday: str, age: str, job: str, position: str, p: [str], doctors: {str}, cabinets: {str}):
    doc = Document()

    section = doc.sections[-1]
    section.page_height = Inches(8.27)
    section.page_width = Inches(5.83)
    section.orientation = WD_ORIENTATION.PORTRAIT

    section.top_margin = Inches(0.5 / 2.54)
    section.bottom_margin = Inches(0.5 / 2.54)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    def add_paragraph_with_spacing(doc, text):
        p = doc.add_paragraph(text)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0

    def set_font_for_table(table, font_name, size):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.line_spacing = 1.0
                    for run in paragraph.runs:
                        run.font.name = font_name
                        run.font.size = size
                        run._element.rPr.rFonts.set(qn('w:ascii'), font_name)
                        run._element.rPr.rFonts.set(qn('w:cs'), font_name)
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                        run._element.rPr.rFonts.set(qn('w:hAnsi'), font_name)

    def set_bold_for_column(table, col_index):
        for row in table.rows:
            cell = row.cells[col_index]
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

    def center_text_vertically(table):
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_paragraph_with_spacing(doc, f'ФИО: {fullName}         Возраст: {age}         Пол: {sex}')
    add_paragraph_with_spacing(doc, f'Дата рождения: {birthday}')
    add_paragraph_with_spacing(doc, f'Место работы: {job}')
    add_paragraph_with_spacing(doc, f'Должность: {position}')
    add_paragraph_with_spacing(doc, f'Вредные факторы: {p}')
    add_paragraph_with_spacing(doc, '')

    table = doc.add_table(rows=len(doctors) + len(cabinets) + 1, cols=3)

    specialties = []

    #Logic
    for name, value in doctors.items():
        specialties.append((name, value))

    for name, value in cabinets.items():
        specialties.append((name, value))


    specialties = sort_specialties(specialties)
    print(specialties)

    for i, (specialty, cabinet) in enumerate(specialties):

        row = table.rows[i]
        row.cells[0].text = specialty
        row.cells[1].text = str(cabinet)
        row.height = Inches(0.2)

    last_row = table.rows[-1]
    last_row.cells[0].merge(last_row.cells[1])
    last_row.cells[0].merge(last_row.cells[2])
    last_row.cells[0].text = "Внимание! После завершения медосмотра маршрутный лист сдайте в регистратуру."

    tbl = table._tbl
    set_font_for_table(table, 'Times New Roman', Pt(10))
    set_bold_for_column(table, 0)
    center_text_vertically(table)

    for cell in tbl.iter_tcs():
        tcPr = cell.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tcBorders.append(border)
        tcPr.append(tcBorders)

    doc.save(f'{fullName}.docx')